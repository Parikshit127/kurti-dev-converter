import os
import tempfile
import unittest
from unittest import mock
import sys

from docx import Document
from docx.oxml.ns import qn

# Fix path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.converter import DocxConverter


def run_ascii_font_name(run):
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return None
    return rpr.rFonts.get(qn('w:ascii'))


class TestDocxConverter(unittest.TestCase):
    def setUp(self):
        self.converter = DocxConverter()

    def test_mixed_language_run_splits_and_fonts_are_selective(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, "mixed.docx")
            output_path = os.path.join(temp_dir, "mixed_out.docx")

            doc = Document()
            para = doc.add_paragraph()
            para.add_run("Hello हिंदी World")
            doc.save(input_path)

            self.converter.convert_file(input_path, output_path)

            converted = Document(output_path)
            runs = converted.paragraphs[0].runs
            text_by_run = [run.text for run in runs]

            self.assertIn("Hello ", text_by_run)
            self.assertIn("fganh", text_by_run)
            self.assertIn(" World", text_by_run)

            hindi_run = next(run for run in runs if run.text == "fganh")
            self.assertEqual(run_ascii_font_name(hindi_run), "Kruti Dev 010")

            for run in runs:
                if run.text in ("Hello ", " World"):
                    self.assertNotEqual(run_ascii_font_name(run), "Kruti Dev 010")

    def test_headers_footers_and_nested_tables_are_processed(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, "stories.docx")
            output_path = os.path.join(temp_dir, "stories_out.docx")

            doc = Document()
            doc.add_paragraph("मुख्य कर्म")

            header_para = doc.sections[0].header.paragraphs[0]
            header_para.text = "कर्म"

            footer_para = doc.sections[0].footer.paragraphs[0]
            footer_para.text = "धर्म"

            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "तालिका"
            nested = cell.add_table(rows=1, cols=1)
            nested.cell(0, 0).text = "कर्म"

            doc.save(input_path)
            self.converter.convert_file(input_path, output_path)

            converted = Document(output_path)
            header_text = "".join(p.text for p in converted.sections[0].header.paragraphs)
            footer_text = "".join(p.text for p in converted.sections[0].footer.paragraphs)
            body_text = "".join(p.text for p in converted.paragraphs)

            self.assertIn("deZ", header_text)
            self.assertIn("/keZ", footer_text)
            self.assertIn("eq[;", body_text)

            table_text = "".join(cell.text for table in converted.tables for row in table.rows for cell in row.cells)
            self.assertIn("rkfydk", table_text)

            nested_converted_text = converted.tables[0].cell(0, 0).tables[0].cell(0, 0).text
            self.assertIn("deZ", nested_converted_text)

    @mock.patch("core.converter.shutil.which", return_value=None)
    def test_doc_conversion_without_libreoffice_fails_with_clear_error(self, _which_mock):
        with self.assertRaises(RuntimeError) as context:
            self.converter.convert_doc_to_docx("/tmp/sample.doc")
        self.assertIn("LibreOffice", str(context.exception))


if __name__ == "__main__":
    unittest.main()
