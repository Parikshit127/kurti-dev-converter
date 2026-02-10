
from docx import Document
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.run import CT_R
from docx.text.run import Run
from .reorder import ReorderEngine
import gc
import logging
import os
import shutil
import subprocess
import tempfile

logger = logging.getLogger(__name__)

class DocxConverter:
    def __init__(self):
        self.engine = ReorderEngine()

    def convert_file(self, input_path, output_path):
        """Convert a document with memory-optimized processing."""
        logger.info(f"Loading document: {input_path}")

        prepared_input_path = input_path
        temp_dir_to_cleanup = None
        doc = None
        try:
            try:
                prepared_input_path, temp_dir_to_cleanup = self.prepare_input_docx(input_path)
            except Exception as e:
                logger.error(f"Failed to prepare input document: {e}")
                raise

            try:
                doc = Document(prepared_input_path)
            except Exception as e:
                logger.error(f"Failed to load document: {e}")
                raise

            logger.info("Document loaded. Starting story traversal.")
            self.process_document(doc)

            # Force garbage collection before saving
            gc.collect()

            logger.info(f"Saving document to: {output_path}")
            try:
                doc.save(output_path)
            except Exception as e:
                logger.error(f"Failed to save document: {e}")
                raise

            logger.info("Document saved successfully")
        finally:
            if temp_dir_to_cleanup:
                shutil.rmtree(temp_dir_to_cleanup, ignore_errors=True)
            if doc is not None:
                del doc
                gc.collect()

    def prepare_input_docx(self, input_path):
        """
        Normalize input to DOCX.

        Returns:
            (docx_path, temporary_directory_to_cleanup_or_none)
        """
        lowered = input_path.lower()
        if lowered.endswith(".docx"):
            return input_path, None
        if lowered.endswith(".doc"):
            return self.convert_doc_to_docx(input_path)
        raise ValueError("Only .doc and .docx files are supported")

    def convert_doc_to_docx(self, input_path):
        """
        Convert legacy .doc to .docx using LibreOffice (soffice).
        """
        soffice = shutil.which("soffice")
        if not soffice:
            raise RuntimeError(
                "Legacy .doc conversion requires LibreOffice ('soffice'), but it is not installed on the server."
            )

        output_dir = tempfile.mkdtemp(prefix="fontchanger_doc_")
        command = [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            output_dir,
            input_path,
        ]
        result = subprocess.run(command, capture_output=True, text=True)
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed (code {result.returncode}): {result.stderr.strip() or result.stdout.strip()}"
            )

        base_name = os.path.splitext(os.path.basename(input_path))[0]
        converted_path = os.path.join(output_dir, f"{base_name}.docx")
        if not os.path.exists(converted_path):
            generated_docx = [
                os.path.join(output_dir, name)
                for name in os.listdir(output_dir)
                if name.lower().endswith(".docx")
            ]
            if len(generated_docx) == 1:
                converted_path = generated_docx[0]
            else:
                raise RuntimeError("LibreOffice conversion did not produce a .docx file.")

        logger.info("Converted legacy .doc to .docx via LibreOffice")
        return converted_path, output_dir

    def process_document(self, doc):
        """Process body + header/footer stories and nested tables."""
        for story_name, story in self.iter_stories(doc):
            try:
                processed = self.process_story(story)
                logger.info(
                    "Processed story '%s' (paragraphs=%s, tables=%s)",
                    story_name,
                    processed["paragraphs"],
                    processed["tables"],
                )
            except Exception as e:
                logger.warning(f"Error processing story '{story_name}': {e}")

            gc.collect()

    def iter_stories(self, doc):
        """Yield all unique text stories in the document."""
        yield "body", doc

        seen_parts = set()
        for section_index, section in enumerate(doc.sections, start=1):
            story_specs = [
                ("header", section.header),
                ("first_page_header", section.first_page_header),
                ("even_page_header", section.even_page_header),
                ("footer", section.footer),
                ("first_page_footer", section.first_page_footer),
                ("even_page_footer", section.even_page_footer),
            ]
            for story_type, story in story_specs:
                partname = str(story.part.partname)
                if partname in seen_parts:
                    continue
                seen_parts.add(partname)
                yield f"section_{section_index}_{story_type}", story

    def process_story(self, story):
        """Process all paragraphs and tables inside a story container."""
        processed_paragraphs = 0
        processed_tables = 0

        for para in story.paragraphs:
            try:
                self.process_paragraph(para)
                processed_paragraphs += 1
            except Exception as e:
                logger.warning(f"Error processing paragraph: {e}")

        for table in story.tables:
            processed_tables += self.process_table(table)

        return {"paragraphs": processed_paragraphs, "tables": processed_tables}

    def process_table(self, table):
        """Recursively process a table, handling nested tables safely."""
        processed_tables = 1
        seen_cells = set()

        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)

                try:
                    story_result = self.process_story(cell)
                    processed_tables += story_result["tables"]
                except Exception as e:
                    logger.warning(f"Error processing table cell: {e}")

        return processed_tables

    def process_paragraph(self, para):
        """Process a single paragraph - iterate over runs and hyperlinks."""
        runs = []

        # Iterate over all children to catch Runs AND Hyperlinks
        for item in para._element:
            if isinstance(item, CT_R):  # It's a normal Run
                runs.append(Run(item, para))
            elif item.tag.endswith('hyperlink'):  # It's a Hyperlink, which contains runs
                for child in item:
                    if isinstance(child, CT_R):
                        runs.append(Run(child, para))

        for run in runs:
            self.process_run(run)

    def process_run(self, run):
        """Process a single run with mixed-language safe segment splitting."""
        text = run.text
        if not text:
            return

        if not text.strip():
            return  # Don't touch empty/whitespace runs

        if not self.contains_devanagari(text):
            return

        try:
            segments = self.engine.process_segments(text)
        except Exception as e:
            # Log but don't crash on individual run errors
            logger.debug(f"Conversion skipped for run due to error: {e}")
            return

        if not segments:
            return

        needs_kruti = any(segment_needs_kruti for _, segment_needs_kruti in segments)
        if not needs_kruti:
            return

        if len(segments) == 1 and segments[0][1]:
            run.text = segments[0][0]
            self.set_kruti_font(run)
            return

        if self.split_run_with_segments(run, segments):
            return

        # Fallback for complex runs: keep legacy behavior for safety.
        run.text = ''.join(segment_text for segment_text, _ in segments)
        self.set_kruti_font(run)

    def split_run_with_segments(self, run, segments):
        """
        Split a run into multiple runs based on Kruti/non-Kruti segments.
        Returns True if split succeeded.
        """
        run_element = run._element
        if not self.is_simple_text_run(run_element):
            return False

        parent = run_element.getparent()
        if parent is None:
            return False

        insert_idx = parent.index(run_element)
        for segment_text, segment_needs_kruti in segments:
            if not segment_text:
                continue

            new_run_element = self.clone_run_with_text(run_element, segment_text)
            parent.insert(insert_idx, new_run_element)
            insert_idx += 1

            if segment_needs_kruti:
                self.set_kruti_font_on_element(new_run_element)

        parent.remove(run_element)
        return True

    def is_simple_text_run(self, run_element):
        """
        Return True only when run contains plain text nodes that are safe to split.
        """
        text_tag = qn('w:t')
        rpr_tag = qn('w:rPr')
        allowed_tags = {text_tag, rpr_tag}

        text_node_count = 0
        for child in run_element:
            if child.tag not in allowed_tags:
                return False
            if child.tag == text_tag:
                text_node_count += 1

        return text_node_count >= 1

    def clone_run_with_text(self, template_run_element, text):
        """Clone run formatting and set new plain text payload."""
        new_run = deepcopy(template_run_element)
        rpr_tag = qn('w:rPr')

        for child in list(new_run):
            if child.tag != rpr_tag:
                new_run.remove(child)

        text_element = OxmlElement('w:t')
        text_element.text = text
        if text != text.strip():
            text_element.set(qn('xml:space'), 'preserve')
        new_run.append(text_element)
        return new_run

    def contains_devanagari(self, text):
        """Check whether text contains Devanagari Unicode characters."""
        return any('\u0900' <= char <= '\u097F' for char in text)

    def set_kruti_font(self, run):
        """Set the Kruti Dev 010 font on a run."""
        try:
            run.font.name = 'Kruti Dev 010'
            # Force XML properties for all font types (important for Word)
            rPr = run._element.get_or_add_rPr()
            fonts = rPr.get_or_add_rFonts()
            
            fonts.set(qn('w:ascii'), 'Kruti Dev 010')
            fonts.set(qn('w:hAnsi'), 'Kruti Dev 010')
            fonts.set(qn('w:cs'), 'Kruti Dev 010')
            fonts.set(qn('w:eastAsia'), 'Kruti Dev 010')
        except Exception as e:
            logger.debug(f"Failed to set font: {e}")

    def set_kruti_font_on_element(self, run_element):
        """Set Kruti Dev font on a raw run XML element."""
        try:
            rpr = run_element.get_or_add_rPr()
            fonts = rpr.get_or_add_rFonts()
            fonts.set(qn('w:ascii'), 'Kruti Dev 010')
            fonts.set(qn('w:hAnsi'), 'Kruti Dev 010')
            fonts.set(qn('w:cs'), 'Kruti Dev 010')
            fonts.set(qn('w:eastAsia'), 'Kruti Dev 010')
        except Exception as e:
            logger.debug(f"Failed to set XML font properties: {e}")
